from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = r"docs\Diseno_UX_UI_y_Sistema_de_Diseno_EL_VITRAL_Completado.docx"

BLUE = "2E74B5"; DARK = "1F4D78"; LIGHT = "E8EEF5"; GRAY = "F2F4F7"; RED = "9B1C1C"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        n=tcMar.find(qn(f'w:{m}'))
        if n is None: n=OxmlElement(f'w:{m}'); tcMar.append(n)
        n.set(qn('w:w'),str(v)); n.set(qn('w:type'),'dxa')

def set_width(cell, dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(dxa)); tcW.set(qn('w:type'),'dxa')

def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style='Table Grid'; t.autofit=False
    if widths is None: widths=[9360//len(headers)]*len(headers)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,LIGHT); set_width(c,widths[i]); set_cell_margins(c)
        for r in c.paragraphs[0].runs: r.bold=True; r.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); set_width(cells[i],widths[i]); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.size=Pt(8.5)
    doc.add_paragraph()
    return t

def p(doc, text='', bold=False, italic=False, color=None, size=None, style=None):
    x=doc.add_paragraph(style=style) if style else doc.add_paragraph()
    r=x.add_run(text); r.bold=bold; r.italic=italic
    if color: r.font.color.rgb=RGBColor.from_string(color)
    if size: r.font.size=Pt(size)
    return x

def bullets(doc, items):
    for item in items: doc.add_paragraph(item, style='List Bullet')

def heading(doc, text, level=1): doc.add_heading(text, level=level)

def page(doc): doc.add_page_break()

def setup(doc):
    sec=doc.sections[0]; sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.8); sec.right_margin=Inches(0.8)
    sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
    normal=doc.styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(10); normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.12
    for name,size,color in [('Title',24,DARK),('Heading 1',16,BLUE),('Heading 2',13,BLUE),('Heading 3',11,DARK)]:
        st=doc.styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=header.add_run('SENA - ADSO | Diseño UX/UI y Sistema de Diseño | EL VITRAL'); r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('Documento de experiencia de usuario e interfaz | Versión 1.1 | 10/08/2026'); r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)

doc=Document(); setup(doc)

# Cover
x=doc.add_paragraph(); x.alignment=WD_ALIGN_PARAGRAPH.CENTER; x.paragraph_format.space_before=Pt(80)
r=x.add_run('SERVICIO NACIONAL DE APRENDIZAJE - SENA'); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(DARK)
x=doc.add_paragraph(); x.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=x.add_run('TECNOLOGÍA EN ANÁLISIS Y DESARROLLO DE SOFTWARE - ADSO'); r.bold=True; r.font.size=Pt(11)
x=doc.add_paragraph(); x.alignment=WD_ALIGN_PARAGRAPH.CENTER; x.paragraph_format.space_before=Pt(36); r=x.add_run('DISEÑO UX/UI Y\nSISTEMA DE DISEÑO'); r.bold=True; r.font.size=Pt(27); r.font.color.rgb=RGBColor.from_string(BLUE)
x=doc.add_paragraph(); x.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=x.add_run('Sistema digital para gestión comercial y operativa EL VITRAL'); r.italic=True; r.font.size=Pt(14)
p(doc,'Documento completado a partir de los requisitos, mockups y código del proyecto. Los datos marcados "POR CONFIRMAR" deben validarse antes de aprobación.',italic=True,color=RED)
table(doc,['Campo','Información'],[
['Nombre del proyecto','Sistema digital para gestión comercial y operativa EL VITRAL'],['Centro de formación','POR CONFIRMAR en SOFIA Plus / datos oficiales de la ficha'],['Regional','Antioquia (confirmar denominación institucional)'],['Ficha','3144585'],['Programa','Tecnología en Análisis y Desarrollo de Software - ADSO'],['Equipo responsable','Simón Sierra López, José Manuel Ruiz Zapata y Juan José Giraldo Monsalve'],['Responsable UX/UI','Equipo EL VITRAL (asignar responsable individual si el instructor lo exige)'],['Product Owner','Simón Sierra López'],['Instructor','POR CONFIRMAR: documentos existentes mencionan a Juan Carlos Quintero Romero y Paula Andrea Bolívar Uribe'],['Versión / fecha','1.1 / 10 de agosto de 2026']], [2700,6660])
page(doc)

heading(doc,'Control documental',1)
table(doc,['Elemento','Detalle'],[['Código documental','ADSO-UXUI-01'],['Nombre','Diseño UX/UI y Sistema de Diseño'],['Clasificación','Documento de diseño, validación y especificación visual'],['Estado','En revisión'],['Propietario','Equipo EL VITRAL / responsable UX/UI'],['Revisor funcional','Cliente o Product Owner'],['Revisor técnico','Líder frontend'],['Revisor académico','Instructor responsable'],['Ubicación oficial','/docs/diseno/ y repositorio del proyecto'],['Periodicidad','En cada iteración que modifique interfaz, flujo o componente']], [2800,6560])
heading(doc,'Historial de versiones',2)
table(doc,['Versión','Fecha','Descripción','Responsable','Aprobado por'],[['0.1','10/04/2026','Wireframes y mockups iniciales','Equipo EL VITRAL','Pendiente'],['1.0','POR CONFIRMAR','Diseño aprobado para implementación','Equipo EL VITRAL','Pendiente'],['1.1','10/08/2026','Inventario ampliado de pantallas, flujos, estados, accesibilidad y trazabilidad','Equipo EL VITRAL','Pendiente']], [900,1300,4050,1700,1410])
heading(doc,'Registro de aprobaciones',2)
table(doc,['Rol','Nombre','Decisión','Fecha','Observaciones'],[['Responsable funcional','POR CONFIRMAR','Pendiente de aprobación','-','Revisar flujos de cotización y pedido'],['Usuario representante','POR CONFIRMAR','Pendiente de validación','-','Realizar prueba de usabilidad'],['Líder frontend','Equipo EL VITRAL','Pendiente de revisión','-','Verificar correspondencia código-diseño'],['Instructor','POR CONFIRMAR','Pendiente de revisión','-','Validación académica']], [1900,2100,2000,1200,2160])
page(doc)

heading(doc,'Tabla de contenido',1)
p(doc,'1. Propósito, alcance y principios\n2. Relación con otros artefactos\n3. Investigación y comprensión de usuarios\n4. Arquitectura de información\n5. Flujos de usuario y navegación\n6. Inventario y especificación de pantallas\n7. Wireframes\n8. Mockups y prototipos\n9. Estados y comportamiento de interfaz\n10. Diseño responsive y compatibilidad\n11. Accesibilidad\n12. Sistema de diseño\n13. Componentes y patrones\n14. UX writing y contenido\n15. Pruebas de usabilidad y validación\n16. Entrega a desarrollo y control de cambios\n17. Trazabilidad, indicadores y aprobación\nAnexos')
page(doc)

heading(doc,'1. Propósito, alcance y principios',1)
p(doc,'Este documento define la experiencia de usuario, arquitectura de información, flujos, pantallas, estados y reglas visuales de EL VITRAL. Su propósito es convertir procesos manuales de cotización, pedidos, inventario y atención al cliente en una experiencia web comprensible, consistente, trazable y verificable.')
heading(doc,'1.1 Objetivos específicos',2)
bullets(doc,['Identificar las necesidades del cliente final y del administrador del negocio.','Diseñar rutas y pantallas que soporten cotización, pedidos, inventario, agenda y administración.','Documentar reglas de acceso para visitante, usuario autenticado y administrador.','Definir componentes reutilizables, estados de interfaz y criterios responsive.','Mantener trazabilidad entre historias, requisitos, diseño, código y pruebas.','Validar flujos críticos con usuarios representantes antes de aprobación final.'])
heading(doc,'1.2 Principios de diseño',2)
table(doc,['Principio','Aplicación en EL VITRAL'],[['Centrado en personas','La cotización reduce cálculos manuales y permite ingresar medidas y cantidades.'],['Consistencia','Tablas, formularios, tarjetas, colores y navegación se reutilizan entre módulos.'],['Claridad','Cada pantalla muestra título, acción principal, estado de carga o mensaje de resultado.'],['Prevención de errores','Validaciones de campos, confirmación al eliminar y regla de fecha antes de entregar un pedido.'],['Accesibilidad','Se deben verificar contraste, foco, etiquetas persistentes y operación con teclado.'],['Simplicidad responsable','Solo se incluyen funciones del alcance: no pagos en línea, facturación DIAN ni modo offline.']], [2300,7060])
page(doc)

heading(doc,'2. Relación con otros artefactos',1)
table(doc,['Artefacto','Aporta a UX/UI','Recibe de UX/UI'],[['Acta de constitución','Problema, objetivo, alcance, equipo y stakeholders.','Evidencia de solución propuesta.'],['Documento de alcance','Límites, perfiles, RF y RNF.','Pantallas, flujos y restricciones de interfaz.'],['ELVITRAL-RF','Historias, roles, criterios y reglas de negocio.','Referencia de UI, flujos y criterios de aceptación.'],['Mockup-ELVITRAL','Diseños visuales y 16 pantallas iniciales.','Versión, enlaces y cambios requeridos.'],['Código frontend/backend','Pantallas, roles, endpoints y estados realmente implementados.','Guía para correspondencia UI-código.'],['Pruebas Jest','Cobertura de autenticación, catálogo, cotización, pedidos, inventario y productos.','Escenarios de validación y evidencia.']], [2200,3550,3610])
p(doc,'Nota de trazabilidad: este documento no reemplaza requisitos, pruebas ni el prototipo editable. Los complementa y debe actualizarse si cambia la interfaz implementada.',italic=True,color=RED)
page(doc)

heading(doc,'3. Investigación y comprensión de usuarios',1)
heading(doc,'3.1 Fuentes de información',2)
table(doc,['Fuente','Propósito','Evidencia / estado'],[['Acta y alcance del proyecto','Comprender el problema de procesos manuales.','Disponible: acta de constitución y DOCAL.'],['Requisitos e historias','Identificar tareas, reglas y roles.','Disponible: ELVITRAL-RF.'],['Análisis del producto implementado','Contrastar diseño con funciones reales.','Disponible: frontend, backend y pruebas.'],['Entrevistas / pruebas con usuarios','Validar lenguaje, fricciones y tareas.','PENDIENTE: no se encontró evidencia formal.']], [2450,3450,3460])
heading(doc,'3.2 Perfiles o personas',2)
table(doc,['ID','Perfil','Objetivos / tareas','Contexto y limitaciones','Fuente'],[['PER-001','Cliente final','Explorar catálogo, cotizar, consultar pedidos, agendar y calificar servicio.','Usa computador o móvil; busca respuesta rápida; no requiere conocimiento técnico.','Alcance y requisitos.'],['PER-002','Administrador','Gestionar usuarios, productos, inventario, cotizaciones, pedidos y agenda.','Usa principalmente escritorio; requiere control y visualización de datos.','Código y alcance.'],['PER-003','Asesor comercial','Atender cotizaciones y hacer seguimiento.','Perfil definido en alcance, pero sin rol separado implementado.','Alcance; pendiente de validar.'],['PER-004','Instalador','Consultar pedidos e instalaciones.','Perfil definido en alcance, pero no implementado.','Alcance; fuera de versión actual.']], [850,1350,3200,2350,1610])
heading(doc,'3.3 Mapa de necesidades',2)
table(doc,['ID','Perfil','Necesidad','Problema actual','Impacto','Prioridad'],[['NEC-001','PER-001','Cotizar con medidas y cantidad.','Proceso manual, lento y con errores.','Alto','Alta'],['NEC-002','PER-001','Consultar estado de cotización y pedido.','Falta de seguimiento centralizado.','Alto','Alta'],['NEC-003','PER-002','Gestionar productos e inventario.','Control visual y pérdida de información.','Alto','Alta'],['NEC-004','PER-002','Controlar agenda y entregas.','Coordinación por llamadas.','Medio','Media'],['NEC-005','PER-002','Conocer satisfacción posterior a entrega.','No hay registro sistemático.','Medio','Media']], [900,1050,2700,2600,1100,1010])
page(doc)

heading(doc,'4. Arquitectura de información',1)
heading(doc,'4.1 Inventario de contenidos y funciones',2)
table(doc,['ID','Contenido / función','Usuario','Prioridad','Ubicación','Requisito'],[['CONT-001','Inicio, proyectos y datos de contacto','Todos','Alta','/','HU-01'],['CONT-002','Catálogo y filtro por tipo','Todos','Alta','/catalogo','RF catálogo'],['CONT-003','Cotización con cálculo','Autenticado','Alta','/cotizar','RF-001, RF-002'],['CONT-004','Autenticación y perfil','Visitante/usuario','Alta','/login, /registro, /perfil','HU-02, RF-004'],['CONT-005','Cotizaciones y pedidos','Usuario/admin','Alta','/cotizaciones, /mis-pedidos','RF pedidos'],['CONT-006','Administración e inventario','Admin','Alta','/admin/*','RF-005'],['CONT-007','Agenda y encuesta','Usuario/admin','Media','Agenda y pedidos','RF-006, RNF-003']], [850,2300,1550,900,2000,1760])
heading(doc,'4.2 Estructura de navegación y 4.3 Mapa del sitio',2)
p(doc,'Inicio > Catálogo > Cotizar; Inicio > Proyectos > Detalle; menú de usuario > Perfil, Mis cotizaciones, Mis pedidos; Administración > Usuarios, Productos, Inventario, Cotizaciones, Pedidos y Agenda. El visitante también accede a Login, Registro y recuperación de contraseña.')
heading(doc,'4.4 Criterios de calidad',2)
bullets(doc,['Etiquetas en español y vocabulario del negocio: cotización, pedido, inventario, entrega.','La ruta activa y el título de página deben orientar al usuario.','Los permisos se validan en interfaz, middleware y API; ocultar botones no sustituye la autorización.','Las tablas deben ofrecer desplazamiento horizontal en móvil.','No debe existir una pantalla sin historia, requisito, necesidad o función trazable.'])
page(doc)

heading(doc,'5. Flujos de usuario y navegación',1)
heading(doc,'5.1 Catálogo de flujos',2)
table(doc,['ID','Nombre','Actor','Objetivo','Inicio / fin','Prioridad'],[['FLU-001','Registro e inicio de sesión','Visitante','Crear cuenta e ingresar.','Registro/Login > Inicio','Alta'],['FLU-002','Restablecer contraseña','Visitante','Recuperar acceso.','Olvidé contraseña > Login','Media'],['FLU-003','Crear cotización','Cliente','Calcular y generar cotización.','Cotizar > Código generado','Alta'],['FLU-004','Consultar y convertir cotización','Cliente/Admin','Ver detalle y crear pedido.','Cotizaciones > Pedido','Alta'],['FLU-005','Consultar pedido y encuesta','Cliente','Ver pedido y evaluar servicio.','Mis pedidos > Encuesta enviada','Media'],['FLU-006','Gestionar cita','Cliente/Admin','Crear o eliminar una cita.','Agenda > Agenda actualizada','Media'],['FLU-007','Gestionar productos e inventario','Admin','Mantener catálogo y stock.','Admin > Registro actualizado','Alta'],['FLU-008','Gestionar pedidos','Admin','Cambiar estado y fecha de entrega.','Admin pedidos > Pedido actualizado','Alta']], [850,1800,1050,2100,2450,1110])
heading(doc,'5.2 Fichas resumidas de flujo',2)
for title,text in [
('FLU-003 - Crear cotización','Actor: cliente autenticado. Disparador: selecciona “Cotizar”. Precondición: sesión activa y productos disponibles. Camino feliz: ingresa datos, selecciona producto, medidas/cantidad, agrega ítems y genera la cotización. Errores: datos incompletos, correo inválido, lista vacía o producto inexistente. Éxito: se crea código único y se presenta confirmación.'),
('FLU-007 - Gestionar productos e inventario','Actor: administrador. Precondición: rol admin validado. Camino feliz: consulta listado, crea o edita producto, activa/desactiva y registra entradas de inventario. Errores: campos requeridos o cantidad inválida. Éxito: tabla y stock se actualizan.'),
('FLU-008 - Gestionar pedidos','Actor: administrador. Camino feliz: consulta pedidos agrupados por estado, selecciona cambio, confirma la acción y registra fecha de entrega para el estado entregado. Regla: no se permite entregar sin fecha de entrega definida.')]:
    heading(doc,title,3); p(doc,text)
heading(doc,'5.3 Reglas de navegación',2)
table(doc,['ID','Origen','Acción','Destino','Condición'],[['NAV-001','Catálogo','Cotizar producto','/cotizar','Conserva producto seleccionado cuando aplica.'],['NAV-002','Cotizar','Generar cotización','Confirmación','Usuario autenticado y datos válidos.'],['NAV-003','Mis cotizaciones','Convertir a pedido','Mis pedidos','Cotización no convertida.'],['NAV-004','Admin','Abrir módulo','/admin/*','Rol admin.'],['NAV-005','Ruta protegida','Sesión no válida','/login','Redirigir sin exponer datos.']], [850,1600,1800,1900,3210])
page(doc)

heading(doc,'6. Inventario y especificación de pantallas',1)
heading(doc,'6.1 Inventario maestro',2)
screens=[['UI-001','Inicio','Página','Todos','Explorar servicios','HU-01','Implementado'],['UI-002','Catálogo','Página','Todos','Consultar productos','RF catálogo','Implementado'],['UI-003','Cotizar','Página','Autenticado','Crear cotización','RF-001/RF-002','Implementado'],['UI-004','Login','Página','Visitante','Iniciar sesión','HU-02','Implementado'],['UI-005','Registro','Página','Visitante','Crear cuenta','HU-02','Implementado'],['UI-006','Recuperar/restablecer contraseña','Página','Visitante','Recuperar acceso','HU-02','Implementado'],['UI-007','Sobre nosotros','Página','Todos','Conocer empresa','HU-01','Implementado'],['UI-008','Detalle de proyecto','Página','Todos','Ver proyecto','HU-01','Implementado'],['UI-009','Perfil / editar perfil','Página','Autenticado','Consultar/actualizar datos','HU-02','Implementado'],['UI-010','Mis cotizaciones','Página/modal','Autenticado','Consultar/convertir','RF cotizaciones','Implementado'],['UI-011','Mis pedidos / encuesta','Página/modal','Autenticado','Consultar y calificar','RF pedidos/RNF-003','Implementado'],['UI-012','Panel admin','Página','Admin','Acceder a gestión','RF-004','Implementado'],['UI-013','Admin usuarios','Página','Admin','Aprobar usuarios','RF-004','Implementado'],['UI-014','Admin productos','Página/modal','Admin','Gestionar productos','RF catálogo','Implementado'],['UI-015','Admin inventario','Página','Admin','Registrar entradas','RF-005','Implementado'],['UI-016','Admin cotizaciones','Página/modal','Admin','Consultar/convertir','RF cotizaciones','Implementado'],['UI-017','Admin pedidos','Página/modal','Admin','Gestionar estados','RF pedidos','Implementado'],['UI-018','Agenda','Widget/página','Autenticado/Admin','Gestionar citas','RF-006','Implementado']]
table(doc,['ID','Pantalla','Tipo','Actor','Objetivo','HU/RF','Estado'],screens,[750,1750,1100,1250,2150,1300,1060])
page(doc)

heading(doc,'6.2 Ficha de pantalla - patrón aplicable',2)
p(doc,'Aplicar esta ficha a cada UI-001 a UI-018. Para no duplicar información, se registran a continuación las reglas comunes y las especificaciones de pantallas críticas.')
table(doc,['Campo','Especificación'],[['Usuarios/roles','Según inventario maestro; visitante, autenticado o administrador.'],['Datos mostrados','Productos, datos de perfil, cotizaciones, pedidos, inventario, citas o mensajes según módulo.'],['Acciones','Consultar, crear, editar, confirmar, cancelar, eliminar, filtrar, convertir o actualizar estado.'],['Permisos','Las rutas admin requieren rol admin. Las áreas privadas requieren sesión válida.'],['Estados','Inicial, carga, vacío cuando aplique, éxito, validación, error, permiso y sesión expirada.'],['Responsive','Móvil en una columna; tablas con overflow horizontal; menús colapsables.'],['Accesibilidad','Etiquetas persistentes, foco visible, contraste, semántica y mensajes asociados por verificar.'],['Prototipo','Mockup-ELVITRAL versión 1.0.1; actualizar para UI-006, UI-009, UI-015 y UI-018.']], [2350,7010])
heading(doc,'Especificación crítica: UI-003 Cotizar',3)
p(doc,'Objetivo: permitir al cliente crear una cotización con datos personales, producto, medidas y cantidad. Componentes: inputs, select, lista de ítems, total, botones de agregar/eliminar/generar. Validaciones: campos obligatorios, correo válido, productos válidos y lista no vacía. Navegación: al éxito muestra código de cotización; sin sesión muestra indicación para iniciar sesión.')
heading(doc,'Especificación crítica: UI-017 Admin pedidos',3)
p(doc,'Objetivo: administrar pedidos por estados pendiente, en proceso, listo y entregado. Acciones: cambiar estado, definir fecha de entrega y confirmar cambios. Regla: no se puede marcar entregado sin fecha de entrega. Estados: carga, vacío, confirmación, éxito y error de servidor.')
heading(doc,'6.3 Matriz pantalla-rol-acción',2)
table(doc,['Pantalla','Rol','Consultar','Crear','Editar','Eliminar','Exportar','Observación'],[['UI-002','Todos','Sí','No','No','No','No','Solo productos activos.'],['UI-003','Autenticado','Sí','Sí','Sí','Sí','No','Cotización.'],['UI-010/UI-011','Autenticado','Sí','Pedido desde cotización','No','No','No','Solo información propia.'],['UI-013','Admin','Sí','Aprobar','Sí','Sí','No','Gestión de usuarios.'],['UI-014','Admin','Sí','Sí','Sí','Sí','No','Productos activos/inactivos.'],['UI-015','Admin','Sí','Sí','No','No','No','Entradas de inventario.'],['UI-016/UI-017','Admin','Sí','Sí','Sí','No','No','Cotizaciones y pedidos.']], [1250,1050,950,950,950,950,950,2360])
page(doc)

heading(doc,'7. Wireframes',1)
heading(doc,'7.1 Registro de wireframes',2)
table(doc,['ID','Pantalla','Objetivo','Fidelidad','Versión','Ubicación','Estado'],[['WF-001 a WF-016','UI-001 a UI-016 iniciales','Validar estructura y navegación','Media/alta','1.0.1','docs/Mockup-ELVITRAL.pdf','Documentado'],['WF-017','UI-015 Inventario','Documentar interfaz implementada','Pendiente','1.1','docs/diseno/pantallas/','Pendiente'],['WF-018','UI-018 Agenda','Documentar widget y página admin','Pendiente','1.1','docs/diseno/pantallas/','Pendiente'],['WF-019','UI-006 y UI-009','Contraseña y edición de perfil','Pendiente','1.1','docs/diseno/pantallas/','Pendiente']], [1200,1700,2250,1100,900,1550,660])
heading(doc,'7.2 Contenido mínimo y 7.3 aprobación',2)
bullets(doc,['Cada wireframe debe incluir código de pantalla, tarea principal, jerarquía, acciones, navegación, estados y relación HU/RF.','Aprobación requerida: la estructura soporta la tarea, considera estados vacíos/errores, coincide con el flujo y no agrega alcance.','Evidencia: fecha, versión y nombre de la persona que valida.'])
heading(doc,'8. Mockups y prototipos',1)
heading(doc,'8.1 Registro de mockups',2)
table(doc,['ID','Pantalla','Dispositivo','Versión','Ubicación','Estado'],[['MK-001 a MK-016','Pantallas documentadas en mockup inicial','Escritorio y móvil','1.0.1','docs/Mockup-ELVITRAL.pdf','Implementado / actualizar'],['MK-017 a MK-019','Pantallas implementadas sin mockup consolidado','Escritorio y móvil','1.1','docs/diseno/pantallas/','Pendiente de crear']], [1450,2350,1750,950,1850,1010])
heading(doc,'8.2 Registro de prototipos y 8.3 versionamiento',2)
p(doc,'PRT-001: prototipo navegable de flujos críticos (registro, cotización, pedido y administración). Herramienta/enlace: POR CONFIRMAR. Versión: 1.1. Instrucciones: incluir ruta de prueba y, si se requieren, credenciales de demostración. Todo cambio posterior a aprobación debe registrarse en la sección 16.3.')
page(doc)

heading(doc,'9. Estados y comportamiento de interfaz',1)
heading(doc,'9.1 Estados obligatorios',2)
table(doc,['Estado','Aplicación en EL VITRAL'],[['Inicial','Contexto, título y acción principal de cada pantalla.'],['Carga','Consulta de productos, pedidos, cotizaciones, inventario, perfil y agenda.'],['Vacío','Catálogo, listas de pedidos/cotizaciones o agenda sin registros.'],['Éxito','Cotización creada, producto guardado, cita creada, encuesta enviada o pedido actualizado.'],['Validación','Campos obligatorios, correo inválido, medidas/cantidad incorrectas.'],['Error sistema','Fallo de red/API; conservar datos ingresados y ofrecer reintento.'],['Permiso insuficiente','Acceso no autorizado a administración; redirigir y no exponer información.'],['Sesión expirada','Solicitar ingreso de nuevo antes de continuar.'],['Confirmación','Eliminar producto/cita y modificar estados sensibles.']], [2300,7060])
heading(doc,'9.2 Matriz de estados por pantalla',2)
table(doc,['Pantalla','Carga','Vacío','Éxito','Validación','Error','Permiso','Observación'],[['UI-002 Catálogo','Sí','Sí','N.A.','N.A.','Sí','N.A.','Productos activos.'],['UI-003 Cotizar','Sí','Sí','Sí','Sí','Sí','Sí','No permite cotizar sin sesión.'],['UI-010 Cotizaciones','Sí','Sí','Sí','N.A.','Sí','Sí','Detalle y conversión.'],['UI-011 Pedidos','Sí','Sí','Sí','Sí','Sí','Sí','Encuesta posterior.'],['UI-014 Productos','Sí','Sí','Sí','Sí','Sí','Sí','Confirmar eliminación.'],['UI-015 Inventario','Sí','Sí','Sí','Sí','Sí','Sí','Cantidad válida.'],['UI-017 Pedidos admin','Sí','Sí','Sí','Sí','Sí','Sí','Fecha obligatoria al entregar.'],['UI-018 Agenda','Sí','Sí','Sí','Sí','Sí','Sí','Fecha no anterior a hoy.']], [1550,800,800,800,950,800,950,2710])
heading(doc,'9.3 Mensajes y prevención de errores',2)
p(doc,'Los mensajes no deben ser genéricos. La implementación actual usa alertas y confirmaciones nativas en varias vistas; se recomienda migrarlas a un componente de notificación/modal accesible y consistente.')
page(doc)

heading(doc,'10. Diseño responsive y compatibilidad',1)
heading(doc,'10.1 Dispositivos y breakpoints',2)
table(doc,['Perfil','Rango','Prioridad','Criterio de aceptación'],[['Móvil','320-767 px','Alta','Sin desbordamiento horizontal innecesario; controles táctiles y formularios en una columna.'],['Tableta','768-1023 px','Media','Grid y formularios se ajustan sin perder jerarquía.'],['Escritorio','>=1024 px','Alta','Aprovecha tablas, paneles, grids y navegación completa.']], [1500,1500,1200,5160])
heading(doc,'10.2 Matriz responsive',2)
table(doc,['Componente','Móvil','Tableta','Escritorio','Regla'],[['Navbar','Menú hamburguesa','Navegación compacta','Enlaces visibles','Usar puntos md/l g.'],['Catálogo','1 columna','2 columnas','3 columnas','Tarjetas sin recorte.'],['Cotización','Columnas apiladas','Ajuste por ancho','Dos columnas','Datos y productos legibles.'],['Tablas','Scroll horizontal','Ajuste por ancho','Tabla completa','No ocultar información crítica.'],['Modal','90% ancho y scroll','Ancho adaptativo','Ancho contenido','Foco y cierre accesibles.']], [1600,1700,1700,1700,2660])
heading(doc,'10.3 Compatibilidad',2)
table(doc,['Plataforma','Versión objetivo','Prioridad','Método','Resultado'],[['Chrome','Última estable','Alta','Prueba manual','Pendiente de registrar'],['Edge','Última estable','Media','Prueba manual','Pendiente de registrar'],['Firefox','Última estable','Media','Prueba manual','Pendiente de registrar'],['Android Chrome / iOS Safari','Actual','Alta','Prueba física o emulador','Pendiente de registrar']], [1800,1600,1200,2450,2310])
page(doc)

heading(doc,'11. Accesibilidad',1)
heading(doc,'11.1 Lista mínima de verificación',2)
access=[['1','Contraste de texto, iconos y fondo','Parcial','Paleta documentada; falta medición WCAG.'],['2','Foco visible y orden de teclado','Pendiente','Requiere prueba manual.'],['3','Campos con etiquetas persistentes','Parcial','Algunos formularios dependen de placeholder.'],['4','Controles con nombre accesible','Pendiente','Revisar iconos y botones.'],['5','Información no solo por color','Pendiente','Revisar badges de estado.'],['6','Zoom y texto legible','Pendiente','Probar 200%.'],['7','Alternativa textual para imágenes','Parcial','Logo/productos tienen alt; auditar resto.'],['8','Tablas con estructura comprensible','Pendiente','Revisar encabezados semánticos.'],['9','Foco en modales y menús','Pendiente','No hay evidencia de focus trap.'],['10','Errores identificables y corregibles','Parcial','Existen alertas; mejorar asociación con campo.'],['11','Objetivos táctiles adecuados','Pendiente','Probar en móvil.'],['12','Reducción de movimiento','Pendiente','No se evidenció configuración.']]
table(doc,['N.','Criterio','Estado','Evidencia / acción'],access,[500,3100,1250,4510])
heading(doc,'11.2 Registro de hallazgos',2)
table(doc,['ID','Pantalla/componente','Severidad','Hallazgo','Acción','Estado'],[['ACC-001','Formularios públicos','Alta','Placeholders usados como guía principal en varios campos.','Incluir labels persistentes y mensajes asociados.','Abierto'],['ACC-002','Modales/menús','Alta','No hay evidencia de manejo de foco con teclado.','Implementar foco inicial, cierre Esc y retorno.','Abierto'],['ACC-003','Estados/badges','Media','Falta validar contraste y significado sin color.','Medir contraste y agregar texto/icono.','Abierto'],['ACC-004','Mensajes','Media','Alertas nativas no aseguran anuncio accesible.','Crear componente de alerta accesible.','Abierto']], [850,1750,1000,2400,2350,1010])
page(doc)

heading(doc,'12. Sistema de diseño',1)
heading(doc,'12.1 Principios y alcance',2)
p(doc,'El sistema de diseño se limita a lo necesario para EL VITRAL: tokens de color, tipografía, espaciado, grid, componentes y estados para una web de cotizaciones y administración. Debe evitar copiar una biblioteca completa que el proyecto no utiliza.')
heading(doc,'12.2 y 12.3 Tokens y paleta',2)
table(doc,['Categoría','Token','Valor','Uso / restricción'],[['Color','color.background.primary','#101828','Fondo principal oscuro.'],['Color','color.surface','#1E2939','Tarjetas y modales.'],['Color','color.surface.secondary','#0F172A','Fondo secundario.'],['Color','color.primary.500','#3B82F6','Acción principal; validar contraste.'],['Color','color.text.primary','#FFFFFF','Texto principal sobre fondo oscuro.'],['Color','color.text.secondary','#CBD5E1','Texto secundario.'],['Estado','color.success/warning/error','Tokens Tailwind por consolidar','Éxito, advertencia y error; no usar solo color.'],['Espaciado','space.1 a space.8','4, 8, 12, 16, 24, 32 px','Escala consistente.'],['Radio','radius.md','8 px recomendado','Inputs, tarjetas, modales.'],['Movimiento','motion.fast','150-200 ms','Transiciones breves; respetar reduced motion.']], [1300,2100,1650,4310])
heading(doc,'12.4 Tipografía y 12.5 grid',2)
table(doc,['Token','Fuente/Peso','Tamaño','Uso'],[['font.display','Inter 700','40-48 px','Hero / portada'],['font.heading.lg','Inter 700','30-36 px','Título de página'],['font.heading.md','Inter 600','20-24 px','Sección'],['font.body.md','Inter 400','16 px','Texto general'],['font.label','Inter 500','14-16 px','Etiquetas'],['font.caption','Inter 400','12-14 px','Ayuda secundaria']], [1900,2300,1500,3660])
p(doc,'Grid: contenedor máximo equivalente a max-w-7xl; márgenes horizontales 16 px móvil, 24 px tableta y 32 px escritorio. Breakpoints por contenido: móvil, md (768 px) y lg (1024 px). Hallazgo técnico: el proyecto usa la clase text-primary, pero debe consolidar la definición del token en CSS/Tailwind.',italic=True,color=RED)
page(doc)

heading(doc,'13. Componentes y patrones',1)
heading(doc,'13.1 Catálogo de componentes',2)
table(doc,['ID','Componente','Variantes/estados','Uso','Implementación'],[['CMP-UI-001','Navbar','Desktop/móvil; abierto/cerrado','Navegación y sesión','NavBar.tsx'],['CMP-UI-002','Botón','Primario, secundario, peligro, disabled, loading','Acciones principales','Reutilizar tokens'],['CMP-UI-003','Campo de formulario','Texto, email, password, número, fecha, textarea; error','Captura de datos','Páginas de auth/cotización'],['CMP-UI-004','Select/combobox','Default, focus, error','Tipo de producto y estado','Cotización/productos'],['CMP-UI-005','Tarjeta','Producto, proyecto, KPI','Resumen de contenido','Inicio/catálogo/admin'],['CMP-UI-006','Tabla','Desktop, scroll móvil, vacío, carga','Listados administrativos','Pedidos/cotizaciones/inventario'],['CMP-UI-007','Modal','Detalle, confirmación, éxito, error','Información sensible','Cotizaciones/pedidos'],['CMP-UI-008','Badge de estado','Pedido/cotización/usuario','Estado visible con texto','Vistas privadas'],['CMP-UI-009','Agenda','Crear/eliminar cita','Gestión de fechas','AgendaWidget.tsx']], [1250,1750,2500,1850,2010])
heading(doc,'13.2 Ficha de componente - Botón',2)
p(doc,'Propósito: ejecutar una acción clara. Anatomía: etiqueta descriptiva, icono opcional y estado de carga. Variantes: primario, secundario, peligro. Estados: default, hover, focus visible, active, disabled y loading. Accesibilidad: nombre accesible, foco visible, operación con Enter/Espacio y sin depender solo de color. Regla: usar “Guardar cambios”, “Generar cotización” o “Eliminar producto”; evitar “Aceptar” si es ambiguo.')
heading(doc,'13.3 Patrones funcionales',2)
table(doc,['Patrón','Objetivo','Reglas mínimas'],[['Formulario','Capturar datos y prevenir errores.','Label persistente, ayuda, validación cercana, conservar datos válidos.'],['Tabla y filtros','Explorar registros.','Encabezados, scroll móvil, carga, vacío y acciones identificables.'],['Confirmación sensible','Evitar acciones accidentales.','Objeto y consecuencia claros; acción secundaria cancelar.'],['Notificación','Informar resultado.','Severidad, mensaje accionable, no revelar detalles internos.'],['Carga','Comunicar procesamiento.','Evitar doble envío y conservar contexto.']], [1800,2300,5260])
page(doc)

heading(doc,'14. UX writing y contenido',1)
heading(doc,'14.1 Voz y tono',2)
table(doc,['Atributo','Aplicación','Evitar'],[['Claro','Frases directas y nombres conocidos por el negocio.','Jerga técnica.'],['Respetuoso','Describe el problema sin culpar.','“Usted se equivocó”.'],['Accionable','Indica qué pasó y siguiente paso.','Errores sin orientación.'],['Consistente','Usa siempre cotización, pedido, producto, entrega.','Sinónimos que confundan.'],['Breve','Un mensaje por idea.','Párrafos largos en alertas.']], [1500,4000,3860])
heading(doc,'14.2 Catálogo de mensajes',2)
table(doc,['ID','Contexto','Tipo','Mensaje','Acción'],[['MSG-001','Cotización creada','Éxito','Cotización creada correctamente. Tu código es {codigo}.','Guardar o consultar código.'],['MSG-002','Formulario incompleto','Validación','Completa los campos obligatorios para continuar.','Señalar primer campo inválido.'],['MSG-003','Correo inválido','Validación','Ingresa un correo electrónico válido.','Corregir correo.'],['MSG-004','Sin permiso','Error seguro','No tienes permisos para realizar esta acción.','Volver o iniciar sesión.'],['MSG-005','Sin resultados','Vacío','No hay registros para mostrar todavía.','Ofrecer acción útil.'],['MSG-006','Entrega sin fecha','Validación','Define la fecha de entrega antes de marcar el pedido como entregado.','Abrir selector de fecha.'],['MSG-007','Eliminar cita','Confirmación','¿Deseas eliminar esta cita? Esta acción no se puede deshacer.','Cancelar o eliminar.']], [800,1550,1200,3950,1860])
heading(doc,'14.3 Reglas de contenido',2)
bullets(doc,['Los placeholders no sustituyen etiquetas.','Botones con verbo y objeto: “Crear producto”, “Guardar cambios”.','Las fechas y valores monetarios se muestran de forma consistente en español colombiano.','Las confirmaciones nombran el objeto afectado y la consecuencia.','Los mensajes técnicos se registran en consola, no se muestran al usuario.'])
page(doc)

heading(doc,'15. Pruebas de usabilidad y validación',1)
heading(doc,'15.1 Plan de prueba',2)
table(doc,['Campo','Contenido'],[['ID','USAB-001'],['Objetivo','Verificar que clientes y administrador completen los flujos críticos sin ayuda mayor.'],['Hipótesis','Un cliente puede generar una cotización y consultar un pedido; un admin puede actualizar un pedido e inventario.'],['Perfiles participantes','3 clientes potenciales y 1 administrador o representante autorizado.'],['Modalidad','Moderada, presencial o remota.'],['Métricas','Éxito de tarea, tiempo, errores, dudas, comentarios y satisfacción.'],['Criterio de aprobación','Al menos 80% de participantes completa cada tarea crítica sin ayuda mayor.'],['Estado','PENDIENTE: no se encontraron resultados de pruebas formales.']], [2100,7260])
heading(doc,'15.2 Guion de tareas',2)
table(doc,['ID','Escenario','Tarea','Éxito esperado','Dato observado'],[['TAR-USAB-01','Cliente nuevo','Registrarse e iniciar sesión.','Accede a la cuenta.','Errores y comprensión.'],['TAR-USAB-02','Cliente','Crear cotización de un producto con medidas.','Obtiene código de cotización.','Tiempo, campos dudosos.'],['TAR-USAB-03','Cliente','Consultar pedido y enviar encuesta.','Ve detalle y completa encuesta.','Comprensión de estados.'],['TAR-USAB-04','Admin','Crear producto y registrar inventario.','Producto/entrada visibles.','Navegación y validaciones.'],['TAR-USAB-05','Admin','Cambiar pedido a entregado.','Define fecha y confirma cambio.','Comprensión de regla.']], [1150,1550,2650,2250,1760])
heading(doc,'15.3 Resultados y 15.4 hallazgos',2)
p(doc,'Dejar estas tablas sin datos hasta ejecutar las sesiones. Registrar participantes anónimos (P-01, P-02), fecha, versión probada y evidencia. No debe declararse validado solamente por preguntar “¿te gustó?”.')
table(doc,['ID','Hallazgo','Evidencia','Severidad','Decisión','Estado'],[['UXH-001','POR REGISTRAR tras prueba','Grabación/notas','-','-','Pendiente'],['UXH-002','POR REGISTRAR tras prueba','Grabación/notas','-','-','Pendiente']], [1000,3000,2100,1100,1250,910])
page(doc)

heading(doc,'16. Entrega a desarrollo y control de cambios',1)
heading(doc,'16.1 Paquete de handoff',2)
table(doc,['Elemento','Obligatorio','Ubicación','Verificación'],[['Prototipo versionado','Sí','docs/diseno/ o Figma','Pendiente de enlace'],['Inventario de pantallas','Sí','Sección 6','Incluye UI-001 a UI-018'],['Tokens y estilos','Sí','Sección 12','Consolidar token primary'],['Componentes y estados','Sí','Secciones 9 y 13','Pendiente auditoría a11y'],['Assets exportables','Según aplique','frontend/public y docs/diseno/assets','Verificar licencias/origen'],['Reglas responsive','Sí','Sección 10','Probar en dispositivos'],['Accesibilidad','Sí','Sección 11','Hallazgos abiertos'],['Mensajes y contenido','Sí','Sección 14','Migrar alertas nativas'],['Trazabilidad UI-HU-RF','Sí','Sección 17','Actualizar por versión'],['Hallazgos pendientes','Sí','Secciones 11 y 15','Cerrar antes de aprobación final']], [2400,1000,2800,3160])
heading(doc,'16.2 Registro de decisiones de diseño',2)
table(doc,['ID','Fecha','Contexto','Decisión','Justificación','Impacto'],[['DD-001','10/04/2026','Alcance','Web responsive, no app nativa.','Alcance y recursos disponibles.','Diseño por breakpoints.'],['DD-002','10/04/2026','Roles','Implementar usuario y administrador.','Son los roles presentes en base de datos.','Asesor/instalador quedan pendientes.'],['DD-003','10/04/2026','Funciones','Excluir pagos y facturación DIAN.','Fuera del MVP.','No diseñar pantallas de pago.'],['DD-004','10/08/2026','UX','Documentar UI-001 a UI-018.','El código tiene más vistas que el mockup inicial.','Actualizar prototipos.']], [850,1200,1800,2100,2100,1310])
heading(doc,'16.3 Solicitudes de cambio y 16.4 correspondencia',2)
p(doc,'CHG-UX-001: actualizar mockups para pantallas agregadas después de la versión 1.0.1. Estado: pendiente. Criterio: toda interfaz implementada debe usar componentes/tokens documentados, contemplar estados, coincidir con mensajes aprobados y registrar cualquier desviación.')
page(doc)

heading(doc,'17. Trazabilidad, indicadores y aprobación',1)
heading(doc,'17.1 Matriz resumida',2)
table(doc,['UI / flujo','Perfil','HU / RF','Componentes','Prueba / evidencia','Estado'],[['UI-001/UI-002','Todos','HU-01 / catálogo','Navbar, cards, filtros','Mockup + código','Implementado'],['UI-003 / FLU-003','Cliente','HU-03 / RF-001, RF-002','Formulario, select, lista','cotizacion.test.js','Implementado'],['UI-004/UI-005','Visitante','HU-02 / RF-004','Inputs, captcha','login/register tests','Implementado'],['UI-010/UI-011','Cliente','Pedidos/cotizaciones','Tabla, modal, badge','pedidos/cotizaciones tests','Implementado'],['UI-014/UI-015','Admin','RF catálogo / RF-005','Tabla, formulario, confirmación','productos/inventario tests','Implementado'],['UI-017 / FLU-008','Admin','Gestión de pedidos','Tabla, modal, select','pedidos tests + código','Implementado'],['UI-018 / FLU-006','Cliente/Admin','RF-006','Agenda/formulario','Código; prueba UX pendiente','Implementado']], [1550,1100,1650,1750,2250,1060])
heading(doc,'17.2 Indicadores sugeridos',2)
table(doc,['Indicador','Método','Meta','Frecuencia','Fuente'],[['Cobertura de pantallas','Pantallas trazadas / inventariadas','100%','Por versión','Matriz UI'],['Flujos críticos validados','Flujos aprobados / críticos','100% antes de entrega','Por iteración','Pruebas UX'],['Éxito de tareas','Participantes que completan / total','>=80%','Por prueba','Resultados USAB'],['Hallazgos críticos abiertos','Conteo','0 antes de entrega','Semanal','Registro UX'],['Cobertura de estados','Estados diseñados / aplicables','100%','Por versión','Fichas UI'],['Componentes reutilizados','Uso de componentes / total aplicable','Definir línea base','Por release','Código/diseño']], [1900,2600,1350,1500,2010])
heading(doc,'17.3 Lista de verificación final',2)
checks=['Perfiles y necesidades con fuente o limitación registrada.','Arquitectura coincide con alcance.','Flujos principales, alternativos y de error documentados.','Todas las pantallas tienen código y trazabilidad.','Mockups/prototipos versionados y actualizados.','Estados de carga, vacío, éxito, error y permisos definidos.','Responsive especificado y probado.','Accesibilidad básica verificada.','Sistema de diseño con tokens, componentes y estados.','Mensajes consistentes y accionables.','Prueba de usabilidad realizada o limitación justificada.','Handoff permite implementar sin ambigüedad.','Implementación corresponde con versión aprobada.','Cambios posteriores registrados.']
for i,c in enumerate(checks,1): p(doc,f'{i}. {c}  Estado actual: parcial o pendiente de validación final.')
heading(doc,'17.4 Aprobación',2)
p(doc,'Las aprobaciones se completan únicamente después de revisar prototipos, evidencia de pruebas de usabilidad, accesibilidad y correspondencia de implementación.')
table(doc,['Rol','Nombre','Decisión','Fecha','Firma/evidencia'],[['Responsable funcional','POR CONFIRMAR','Pendiente','-','-'],['Usuario representante','POR CONFIRMAR','Pendiente','-','-'],['Líder frontend','POR CONFIRMAR','Pendiente','-','-'],['Instructor','POR CONFIRMAR','Pendiente','-','-']], [2000,2300,1800,1400,1860])
page(doc)

heading(doc,'Anexo A. Convenciones de identificación',1)
table(doc,['Elemento','Prefijo','Ejemplo'],[['Persona/perfil','PER','PER-001'],['Necesidad','NEC','NEC-001'],['Contenido','CONT','CONT-001'],['Flujo','FLU','FLU-001'],['Navegación','NAV','NAV-001'],['Pantalla','UI','UI-001'],['Wireframe','WF','WF-001'],['Mockup','MK','MK-001'],['Prototipo','PRT','PRT-001'],['Componente','CMP-UI','CMP-UI-001'],['Mensaje','MSG','MSG-001'],['Prueba usabilidad','USAB','USAB-001'],['Hallazgo UX','UXH','UXH-001'],['Decisión diseño','DD','DD-001'],['Cambio UX','CHG-UX','CHG-UX-001']], [3400,2700,3260])
heading(doc,'Anexo B. Estructura recomendada',1)
p(doc,'/docs/diseno/\n  investigacion/\n  arquitectura-informacion/\n  flujos/\n  pantallas/\n  sistema-diseno/\n  pruebas-usabilidad/\n  decisiones/\n  assets/\n  README.md')
heading(doc,'Anexo C. Evidencias por madurez',1)
table(doc,['Nivel','Evidencia para EL VITRAL'],[['Formulación','Perfiles preliminares, mapa de sitio y wireframes de cotización/pedidos.'],['Construcción inicial','Mockups versionados, inventario UI-001 a UI-018, estados y reglas responsive.'],['Producto funcional','Prototipo, tokens, componentes, accesibilidad básica y pruebas de usabilidad.'],['Producto profesional','Biblioteca implementada, pruebas visuales, auditoría de accesibilidad y control de versiones.']], [2600,6760])
heading(doc,'Anexo D. Pendientes que no deben ocultarse',1)
bullets(doc,['Confirmar centro de formación, regional, instructor y aprobadores.','Crear/actualizar mockups para pantallas implementadas fuera del documento 1.0.1.','Ejecutar pruebas de usabilidad y registrar resultados reales.','Medir contraste y probar teclado, foco, zoom, modales y lectores de pantalla.','Consolidar el token primario y reemplazar alertas/confirmaciones nativas por componentes accesibles.','Probar en Chrome, Edge, Firefox, Android y Safari/iOS antes de aprobar compatibilidad.'])
p(doc,'Fuentes revisadas: Acta de constitución, DOCAL, ELVITRAL-RF, Mockup-ELVITRAL, código frontend/backend, esquema de base de datos y suite de pruebas del repositorio.',italic=True)

doc.save(OUT)
print(OUT)
